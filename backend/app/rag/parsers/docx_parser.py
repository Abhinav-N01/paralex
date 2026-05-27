"""
ParaLex — DOCX Parser (Word documents)
Handles prior memos, letters, legal correspondence.
"""
import logging
from pathlib import Path

from docx import Document as DocxDocument

from app.rag.parsers.pdf_parser import ParsedDocument, ParsedPage, _clean_text

logger = logging.getLogger(__name__)

# Approximate chars per "page" when splitting a Word doc into logical pages
CHARS_PER_PAGE = 3000


def parse_docx(file_path: Path) -> ParsedDocument:
    result = ParsedDocument()

    try:
        doc = DocxDocument(str(file_path))

        full_text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                # Preserve heading hierarchy
                if para.style.name.startswith("Heading"):
                    full_text_parts.append(f"\n## {para.text.strip()}\n")
                else:
                    full_text_parts.append(para.text.strip())

        # Extract tables
        for table in doc.tables:
            table_lines = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_lines.append(" | ".join(cells))
            if table_lines:
                full_text_parts.append("\n[TABLE]\n" + "\n".join(table_lines) + "\n[/TABLE]")

        full_text = "\n".join(full_text_parts)
        full_text = _clean_text(full_text)

        # Split into logical pages
        chunks = _split_into_pages(full_text, CHARS_PER_PAGE)
        result.total_pages = len(chunks)

        for i, chunk in enumerate(chunks, start=1):
            result.pages.append(ParsedPage(page_number=i, text=chunk))

    except Exception as e:
        logger.error(f"DOCX parse failed for {file_path}: {e}")
        result.extraction_warnings.append(f"DOCX parse error: {str(e)}")

    return result


def _split_into_pages(text: str, chars_per_page: int) -> list[str]:
    """Split long text into approximate 'pages' at paragraph boundaries."""
    paragraphs = text.split("\n\n")
    pages = []
    current = []
    current_len = 0

    for para in paragraphs:
        if current_len + len(para) > chars_per_page and current:
            pages.append("\n\n".join(current))
            current = [para]
            current_len = len(para)
        else:
            current.append(para)
            current_len += len(para)

    if current:
        pages.append("\n\n".join(current))

    return pages or [text]
